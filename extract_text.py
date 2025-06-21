import spacy
from PyPDF2 import PdfReader
import docx
from docx import Document
import string
import os
import zipfile


def extract_text_debug(path):
    """Extraction de texte avec debugging détaillé"""
    text = ""
    file_ext = os.path.splitext(path)[1].lower()

    print(f"[DEBUG] Fichier: {path}")
    print(f"[DEBUG] Extension: {file_ext}")
    print(f"[DEBUG] Fichier existe: {os.path.exists(path)}")
    print(f"[DEBUG] Taille fichier: {os.path.getsize(path) if os.path.exists(path) else 'N/A'} bytes")

    try:
        if file_ext == '.pdf':
            reader = PdfReader(path)
            print(f"[DEBUG] Nombre de pages PDF: {len(reader.pages)}")
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                print(f"[DEBUG] Page {i + 1}: {len(page_text)} caractères")
                text += page_text

        elif file_ext in ['.doc', '.docx']:
            print(f"[DEBUG] Tentative d'ouverture du fichier Word...")

            # Vérifier si le fichier est un vrai fichier Word
            try:
                with zipfile.ZipFile(path, 'r') as zip_file:
                    file_list = zip_file.namelist()
                    print(f"[DEBUG] Contenu du fichier ZIP: {file_list[:5]}...")  # Premiers 5 fichiers
                    if 'word/document.xml' not in file_list:
                        print("[DEBUG] ERREUR: Pas de document.xml trouvé - fichier Word corrompu")
                        return ""
            except zipfile.BadZipFile:
                print("[DEBUG] ERREUR: Fichier n'est pas un ZIP valide")
                return ""

            # Extraction avec python-docx
            try:
                doc = Document(path)
                print(f"[DEBUG] Document ouvert avec succès")
                print(f"[DEBUG] Nombre de paragraphes: {len(doc.paragraphs)}")

                # Extraction des paragraphes
                paragraph_count = 0
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text += paragraph.text + "\n"
                        paragraph_count += 1

                print(f"[DEBUG] Paragraphes avec contenu: {paragraph_count}")

                # Extraction des tableaux
                print(f"[DEBUG] Nombre de tableaux: {len(doc.tables)}")
                table_text = ""
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                table_text += cell_text + " "

                text += table_text
                print(f"[DEBUG] Texte des tableaux: {len(table_text)} caractères")

            except Exception as docx_error:
                print(f"[DEBUG] Erreur python-docx: {docx_error}")

                # Fallback: essayer avec une méthode alternative
                print("[DEBUG] Tentative avec méthode alternative...")
                try:
                    import xml.etree.ElementTree as ET

                    with zipfile.ZipFile(path) as docx_zip:
                        # Lire le document principal
                        try:
                            doc_xml = docx_zip.read('word/document.xml')
                            root = ET.fromstring(doc_xml)

                            # Extraire tout le texte
                            for elem in root.iter():
                                if elem.text:
                                    text += elem.text + " "

                            print(f"[DEBUG] Extraction XML réussie: {len(text)} caractères")

                        except KeyError:
                            print("[DEBUG] Pas de word/document.xml trouvé")

                except Exception as xml_error:
                    print(f"[DEBUG] Erreur extraction XML: {xml_error}")

        else:
            print(f"[DEBUG] Format de fichier non supporté: {file_ext}")
            return ""

    except Exception as e:
        print(f"[DEBUG] Erreur générale d'extraction: {e}")
        return ""

    print(f"[DEBUG] Texte final extrait: {len(text)} caractères")
    if text:
        print(f"[DEBUG] Aperçu du texte: '{text[:200]}...'")
    else:
        print("[DEBUG] AUCUN TEXTE EXTRAIT!")

    return text


nlp = spacy.load("fr_core_news_lg")
stop_words = nlp.Defaults.stop_words

# Liste des langues courantes
LANGUES_COURANTES = {
    'français', 'anglais', 'espagnol', 'allemand', 'italien', 'portugais',
    'chinois', 'japonais', 'arabe', 'russe', 'hindi', 'néerlandais',
    'suédois', 'norvégien', 'danois', 'finlandais', 'polonais', 'tchèque',
    'hongrois', 'roumain', 'bulgare', 'grec', 'turc', 'hébreu', 'coréen',
    'thaï', 'vietnamien', 'indonésien', 'malais', 'tagalog', 'swahili',
    'french', 'english', 'spanish', 'german', 'italian', 'portuguese',
    'chinese', 'japanese', 'arabic', 'russian', 'dutch', 'swedish',
    'norwegian', 'danish', 'finnish', 'polish', 'czech', 'hungarian',
    'romanian', 'bulgarian', 'greek', 'turkish', 'hebrew', 'korean',
    'thai', 'vietnamese', 'indonesian', 'malay', 'tagalog', 'swahili'
}


def clean_text(t):
    t = t.strip().lower()
    t = t.translate(str.maketrans('', '', string.punctuation))
    return t


def is_valid_entity(ent_text):
    t = clean_text(ent_text)
    return len(t) > 2 and not all(w in stop_words for w in t.split())


def is_language(text):
    """Détecter si un texte correspond à une langue"""
    cleaned = clean_text(text)
    return cleaned in LANGUES_COURANTES


def extract_entities_structured_auto_debug(text):
    """Extraction d'entités avec debugging détaillé"""
    print(f"[DEBUG] Analyse des entités sur {len(text)} caractères")

    if not text or not text.strip():
        print("[DEBUG] Texte vide ou null!")
        return {
            "competences": [],
            "experiences": [],
            "diplomes": [],
            "postes": [],
            "langues": []
        }

    print(f"[DEBUG] Aperçu du texte à analyser: '{text[:100]}...'")

    try:
        doc = nlp(text)
        print(f"[DEBUG] Document spaCy créé avec {len(doc)} tokens")
        print(f"[DEBUG] Entités détectées par spaCy: {len(doc.ents)}")

        # Debug des entités
        for ent in doc.ents:
            print(f"[DEBUG] Entité: '{ent.text}' - Label: {ent.label_}")

    except Exception as e:
        print(f"[DEBUG] Erreur lors de l'analyse spaCy: {e}")
        return {
            "competences": [],
            "experiences": [],
            "diplomes": [],
            "postes": [],
            "langues": []
        }

    competences = set()
    experiences = set()
    diplomes = set()
    postes = set()
    langues = set()

    # Recherche explicite de langues dans le texte
    lines = text.split('\n')
    print(f"[DEBUG] Analyse de {len(lines)} lignes pour les langues")

    for line in lines:
        line_lower = line.lower()
        if any(keyword in line_lower for keyword in ['langue', 'language', 'idioma', 'parlé', 'écrit', 'oral']):
            print(f"[DEBUG] Ligne avec langues trouvée: '{line}'")
            words = line.split()
            for word in words:
                if is_language(word):
                    langues.add(word.capitalize())
                    print(f"[DEBUG] Langue détectée: {word}")

    # Analyse des entités nommées
    entities_processed = 0
    for ent in doc.ents:
        label = ent.label_.lower()
        ent_text = ent.text.strip()

        if not is_valid_entity(ent_text):
            continue

        cleaned = ent_text.strip()
        entities_processed += 1

        print(f"[DEBUG] Traitement entité #{entities_processed}: '{cleaned}' ({label})")

        # Vérifier d'abord si c'est une langue
        if is_language(cleaned):
            langues.add(cleaned.capitalize())
            print(f"[DEBUG] -> Classée comme langue")
        elif label in ["skill", "work_of_art", "product", "misc"]:
            competences.add(cleaned)
            print(f"[DEBUG] -> Classée comme compétence")
        elif label in ["org", "event"] or "stage" in cleaned.lower() or "poste" in cleaned.lower():
            experiences.add(cleaned)
            print(f"[DEBUG] -> Classée comme expérience")
        elif any(keyword in cleaned.lower() for keyword in ["bac", "licence", "master", "bts", "but", "diplôme"]):
            diplomes.add(cleaned)
            print(f"[DEBUG] -> Classée comme diplôme")
        elif any(keyword in cleaned.lower() for keyword in
                 ["développeur", "analyst", "ingénieur", "consultant", "data"]):
            postes.add(cleaned)
            print(f"[DEBUG] -> Classée comme poste")

    # Recherche de motifs linguistiques pour les langues
    import re
    language_patterns = [
        r'(français|anglais|espagnol|allemand|italien|portugais|chinois|japonais|arabe|russe)\s*[:\-]?\s*(courant|fluent|natif|bilingue|débutant|intermédiaire|avancé)?',
        r'(french|english|spanish|german|italian|portuguese|chinese|japanese|arabic|russian)\s*[:\-]?\s*(fluent|native|beginner|intermediate|advanced)?'
    ]

    for pattern in language_patterns:
        matches = re.finditer(pattern, text.lower())
        for match in matches:
            langue = match.group(1)
            langues.add(langue.capitalize())
            print(f"[DEBUG] Langue détectée par regex: {langue}")

    result = {
        "competences": sorted(competences),
        "experiences": sorted(experiences),
        "diplomes": sorted(diplomes),
        "postes": sorted(postes),
        "langues": sorted(langues)
    }

    print(f"[DEBUG] Résultat final:")
    for key, values in result.items():
        print(f"[DEBUG] {key}: {len(values)} éléments - {values}")

    return result


# Fonctions d'alias pour la compatibilité
def extract_text(path):
    return extract_text_debug(path)


def extract_entities_structured_auto(text):
    return extract_entities_structured_auto_debug(text)